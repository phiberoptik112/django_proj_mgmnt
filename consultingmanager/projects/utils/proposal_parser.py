"""
Proposal Parser Utility
Extracts text from .docx proposal files and auto-categorizes scope of work.
"""

import logging
from typing import Dict, Any, List, Optional
from pathlib import Path
import re

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Install it with: pip install python-docx")

logger = logging.getLogger(__name__)

# Predefined scope categories with keywords for matching
SCOPE_CATEGORIES = {
    'Emergency Generator Noise': [
        'emergency generator', 'generator noise', 'gen noise', 'backup generator',
        'diesel generator', 'generator sound', 'generator acoustics', 'standby generator'
    ],
    'Environmental Noise': [
        'environmental noise', 'ambient noise', 'noise survey', 'noise assessment',
        'traffic noise', 'aircraft noise', 'noise monitoring', 'noise measurement'
    ],
    'STC Testing': [
        'stc', 'sound transmission class', 'stc test', 'stc rating', 'stc measurement',
        'partition testing', 'wall testing', 'sound isolation test'
    ],
    'Room Acoustics': [
        'room acoustics', 'acoustic design', 'reverberation', 'rt60', 'rt 60',
        'acoustic treatment', 'acoustic analysis', 'acoustic consultant',
        'speech intelligibility', 'noise criteria', 'nc rating'
    ],
    'Sound Isolation': [
        'sound isolation', 'noise isolation', 'sound barrier', 'acoustic barrier',
        'soundproofing', 'sound control', 'noise control', 'partition isolation'
    ],
    'Mechanical Noise & Vibration': [
        'mechanical noise', 'vibration', 'hvac noise', 'mechanical system',
        'equipment noise', 'm&e noise', 'vibration isolation', 'mechanical vibration'
    ],
    'Acoustic Consulting': [
        'acoustic consulting', 'acoustic services', 'acoustic design services',
        'acoustic engineering', 'acoustic analysis', 'acoustic assessment'
    ],
    'Building Acoustics': [
        'building acoustics', 'architectural acoustics', 'building noise',
        'structure-borne noise', 'impact noise', 'footfall noise'
    ],
}

def _normalize_text(text: str) -> str:
    """Normalize text for keyword matching"""
    return text.lower().strip()

def _calculate_confidence_score(text: str, keywords: List[str]) -> float:
    """
    Calculate confidence score based on keyword matches
    
    Args:
        text: Text to search in
        keywords: List of keywords to match
        
    Returns:
        Confidence score between 0.0 and 1.0
    """
    if not text or not keywords:
        return 0.0
    
    normalized_text = _normalize_text(text)
    matches = 0
    total_keywords = len(keywords)
    
    for keyword in keywords:
        normalized_keyword = _normalize_text(keyword)
        if normalized_keyword in normalized_text:
            matches += 1
    
    # Calculate score: more matches = higher confidence
    # Also consider if keyword appears multiple times
    if matches == 0:
        return 0.0
    
    base_score = matches / total_keywords
    
    # Boost score if multiple keywords match
    if matches > 1:
        base_score = min(1.0, base_score * 1.2)
    
    # Boost score if keyword appears multiple times in text
    keyword_occurrences = sum(normalized_text.count(_normalize_text(kw)) for kw in keywords)
    if keyword_occurrences > matches:
        frequency_boost = min(0.2, (keyword_occurrences - matches) * 0.05)
        base_score = min(1.0, base_score + frequency_boost)
    
    return base_score

class ProposalParser:
    """Parser for .docx proposal files"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = None
        self.raw_text = ""
        self.logger = logging.getLogger(__name__)
    
    def extract_text(self) -> str:
        """
        Extract all text from the .docx file
        
        Returns:
            Extracted text as a single string
        """
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx library not available")
            return ""
        
        try:
            self.doc = Document(self.file_path)
            paragraphs = []
            
            for paragraph in self.doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Also extract text from tables
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            paragraphs.append(cell_text)
            
            self.raw_text = "\n".join(paragraphs)
            self.logger.info(f"Extracted {len(self.raw_text)} characters from {self.file_path}")
            return self.raw_text
        
        except Exception as e:
            self.logger.error(f"Error extracting text from .docx file {self.file_path}: {e}", exc_info=True)
            return ""
    
    def extract_structure(self) -> Dict[str, Any]:
        """
        Extract document structure (headings, sections)
        
        Returns:
            Dict with document structure information
        """
        if not self.doc:
            return {}
        
        structure = {
            'headings': [],
            'sections': [],
        }
        
        try:
            for paragraph in self.doc.paragraphs:
                # Check if paragraph is a heading (has heading style)
                if paragraph.style and 'heading' in paragraph.style.name.lower():
                    heading_text = paragraph.text.strip()
                    if heading_text:
                        heading_level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                        structure['headings'].append({
                            'text': heading_text,
                            'level': heading_level,
                        })
        except Exception as e:
            self.logger.warning(f"Error extracting document structure: {e}")
        
        return structure
    
    def categorize_scope(self, text: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        Auto-categorize scope of work using keyword matching
        
        Args:
            text: Text to categorize. If None, uses self.raw_text
            
        Returns:
            List of detected categories with confidence scores
        """
        if text is None:
            if not self.raw_text:
                self.extract_text()
            text = self.raw_text
        
        if not text:
            return []
        
        detected_categories = []
        
        for category_name, keywords in SCOPE_CATEGORIES.items():
            confidence = _calculate_confidence_score(text, keywords)
            
            if confidence > 0.0:
                detected_categories.append({
                    'category_name': category_name,
                    'confidence_score': confidence,
                    'keywords_matched': [kw for kw in keywords if _normalize_text(kw) in _normalize_text(text)],
                })
        
        # Sort by confidence score (highest first)
        detected_categories.sort(key=lambda x: x['confidence_score'], reverse=True)
        
        self.logger.info(f"Detected {len(detected_categories)} scope categories")
        return detected_categories
    
    def parse(self) -> Dict[str, Any]:
        """
        Parse the proposal document and return complete results
        
        Returns:
            Dict with extracted text, structure, and categorized scope
        """
        result = {
            'raw_text': '',
            'structure': {},
            'categories': [],
            'file_path': self.file_path,
            'file_name': Path(self.file_path).name,
        }
        
        try:
            # Extract text
            result['raw_text'] = self.extract_text()
            
            # Extract structure
            result['structure'] = self.extract_structure()
            
            # Categorize scope
            result['categories'] = self.categorize_scope()
            
            self.logger.info(f"Successfully parsed proposal: {result['file_name']}")
        
        except Exception as e:
            self.logger.error(f"Error parsing proposal {self.file_path}: {e}", exc_info=True)
            result['error'] = str(e)
        
        return result

