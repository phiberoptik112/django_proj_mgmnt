import os
import pandas as pd
import re
from typing import List, Dict
from pathlib import Path
import pdfplumber
from docx import Document
from datetime import datetime
import matplotlib
matplotlib.use('Agg')  # Set the backend to non-interactive Agg
import matplotlib.pyplot as plt
from matplotlib.cm import get_cmap
import warnings
import logging
import networkx as nx

# Configure logging with more detailed format
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Suppress pdfplumber warnings about CropBox
warnings.filterwarnings('ignore', message='CropBox missing from /Page, defaulting to MediaBox')

def extract_text_from_file(file_path: str) -> str:
    """Extract text content from various file types"""
    logger.debug(f"Attempting to extract text from file: {file_path}")
    extension = Path(file_path).suffix.lower()
    logger.debug(f"File extension detected: {extension}")
    
    try:
        if extension == '.txt':
            logger.debug("Processing text file")
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif extension == '.pdf':
            logger.debug("Processing PDF file")
            with pdfplumber.open(file_path) as pdf:
                logger.debug(f"PDF opened successfully. Total pages: {len(pdf.pages)}")
                text = []
                for i, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text() or ''
                        text.append(page_text)
                        logger.debug(f"Successfully extracted text from page {i}/{len(pdf.pages)}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {i} in {file_path}: {str(e)}")
                        text.append('')
                return ' '.join(text)
        elif extension in ['.doc', '.docx']:
            logger.debug("Processing Word document")
            doc = Document(file_path)
            logger.debug(f"Word document opened successfully. Paragraphs: {len(doc.paragraphs)}")
            return ' '.join(paragraph.text for paragraph in doc.paragraphs)
        else:
            logger.warning(f"Unsupported file type: {extension}")
            return ''
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}", exc_info=True)
        return ''

def find_dollar_amounts(text: str) -> List[Dict]:
    """Find dollar amounts associated with 'lump sum fee' in text"""
    # Pattern for finding "lump sum fee" followed by dollar amount
    pattern = r'lump sum fee of (?:[\$]?\d{1,3}(?:,\d{3})*(?:\.\d{2})?)'
    
    # Find all matches with surrounding context
    findings = []
    for match in re.finditer(pattern, text, re.IGNORECASE):
        # Extract just the dollar amount portion
        amount = re.search(r'(?:[\$]?\d{1,3}(?:,\d{3})*(?:\.\d{2})?)', match.group()).group()
        
        # Get surrounding context
        start = max(0, match.start() - 50)
        end = min(len(text), match.end() + 50)
        context = text[start:end].strip()
        
        findings.append({
            'integers_found': amount,
            'sentence': context
        })
    
    return findings

def create_project_db(project_path: str) -> Dict[str, pd.DataFrame]:
    """Create pandas DataFrames of files in project folder structure"""
    logger.info(f"Creating project database for: {project_path}")
    folder_dfs = {}
    
    for root, dirs, files in os.walk(project_path):
        folder_name = os.path.basename(root)
        logger.debug(f"Processing folder: {folder_name}")
        logger.debug(f"Found {len(files)} files and {len(dirs)} subdirectories")
        
        if files:
            file_data = []
            for file in files:
                file_path = os.path.join(root, file)
                file_type = os.path.splitext(file)[1].lower()
                logger.debug(f"Processing file: {file} (type: {file_type})")
                
                file_data.append({
                    'filename': file,
                    'path': file_path,
                    'file_type': file_type
                })
            
            df = pd.DataFrame(file_data)
            logger.debug(f"Created DataFrame for folder {folder_name} with {len(df)} files")
            
            if folder_name in folder_dfs:
                logger.debug(f"Merging with existing data for folder: {folder_name}")
                folder_dfs[folder_name] = pd.concat([folder_dfs[folder_name], df], ignore_index=True)
            else:
                folder_dfs[folder_name] = df
    
    logger.info(f"Project database creation complete. Total folders processed: {len(folder_dfs)}")
    return folder_dfs

def analyze_proposals(project_paths: List[str]) -> pd.DataFrame:
    """
    Analyze proposal documents from multiple projects to find dollar amounts.
    """
    all_findings = []
    
    for path in project_paths:
        # Walk through project directory
        for root, _, files in os.walk(path):
            for file in files:
                if 'proposal' in file.lower():
                    file_path = os.path.join(root, file)
                    text_content = extract_text_from_file(file_path)
                    
                    if text_content:
                        findings = find_dollar_amounts(text_content)
                        for finding in findings:
                            finding['source_file'] = file
                            finding['project_path'] = path
                            all_findings.append(finding)
    
    # Create DataFrame from findings
    if all_findings:
        results_df = pd.DataFrame(all_findings)
    else:
        results_df = pd.DataFrame(columns=['source_file', 'project_path', 'integers_found', 'sentence'])
    
    return results_df

def get_year_project_paths(base_path: str, year: str) -> List[str]:
    """
    Get a list of all project folder paths for a given year.
    
    Args:
        base_path (str): Base path to the projects directory (e.g. "//DLA-04/Shared/KAILUA PROJECTS/")
        year (str): Year to search for (e.g. "2023")
        
    Returns:
        List[str]: List of full paths to all project folders for that year
    """
    year_path = os.path.join(base_path, year)
    project_paths = []
    
    # Check if year directory exists
    if not os.path.exists(year_path):
        print(f"Warning: Year directory {year_path} not found")
        return project_paths
        
    # Get all subdirectories in the year folder
    try:
        for item in os.listdir(year_path):
            full_path = os.path.join(year_path, item)
            if os.path.isdir(full_path):
                # Only include if it starts with YY-### format
                if re.match(r'\d{2}-\d{3}', item):
                    project_paths.append(full_path)
                    
    except PermissionError:
        print(f"Warning: Permission denied accessing {year_path}")
    except Exception as e:
        print(f"Warning: Error accessing {year_path}: {str(e)}")
        
    return sorted(project_paths)

def analyze_scope_of_work(project_paths: List[str]) -> pd.DataFrame:
    """Analyze proposal documents for scope of work"""
    logger.info("Starting scope of work analysis")
    all_scopes = []
    
    for project_path in project_paths:
        logger.debug(f"Processing project: {project_path}")
        for root, _, files in os.walk(project_path):
            for filename in files:
                if filename.lower().endswith(('.doc', '.docx', '.pdf')):
                    file_path = os.path.join(root, filename)
                    logger.debug(f"Analyzing file: {filename}")
                    
                    try:
                        text_content = extract_text_from_file(file_path)
                        if text_content:
                            logger.debug("Successfully extracted text content")
                            scope_match = re.search(r'(?i)scope\s+of\s+work\s*:?(.*?)(?:\n\s*[A-Z][A-Z\s]+:|$)', 
                                                  text_content, re.DOTALL)
                            # Look for Basic Services section
                            basic_services_match = re.search(r'(?i)basic\s+services\s*:?(.*?)(?:\n\s*[A-Z][A-Z\s]+:|$)', 
                                                          text_content, re.DOTALL)
                            additional_services_match = re.search(r'(?i)additional\s+services\s*:?(.*?)(?:\n\s*[A-Z][A-Z\s]+:|$)', 
                                                                  text_content, re.DOTALL)
                            if basic_services_match:
                                logger.debug("Found basic services section")
                                services_text = basic_services_match.group(1).strip()
                                service_items = re.split(r'\n\s*[•\-\d]+\.?\s*', services_text)
                                service_items = [item.strip() for item in service_items if item.strip()]
                                
                                logger.debug(f"Found {len(service_items)} basic service items")
                                for item in service_items:
                                    scope_info = {
                                        'project_path': project_path,
                                        'source_file': filename,
                                        'scope_item': item,
                                        'category': 'basic_services'
                                    }
                                    all_scopes.append(scope_info)
                            else:
                                logger.debug("No basic services section found in document")
                            if additional_services_match:
                                logger.debug("Found additional services section")
                                services_text = additional_services_match.group(1).strip()
                                service_items = re.split(r'\n\s*[•\-\d]+\.?\s*', services_text)
                                service_items = [item.strip() for item in service_items if item.strip()]
                                
                                logger.debug(f"Found {len(service_items)} additional service items")
                                for item in service_items:
                                    scope_info = {
                                        'project_path': project_path,
                                        'source_file': filename,
                                        'scope_item': item,
                                        'category': 'additional_services'
                                    }
                                    all_scopes.append(scope_info)
                            else:
                                logger.debug("No additional services section found in document")
                            if scope_match:
                                logger.debug("Found scope of work section")
                                scope_text = scope_match.group(1).strip()
                                scope_items = re.split(r'\n\s*[•\-\d]+\.?\s*', scope_text)
                                scope_items = [item.strip() for item in scope_items if item.strip()]
                                
                                logger.debug(f"Found {len(scope_items)} scope items")
                                for item in scope_items:
                                    category = categorize_scope_item(item)
                                    logger.debug(f"Categorized scope item as: {category}")
                                    scope_info = {
                                        'project_path': project_path,
                                        'source_file': filename,
                                        'scope_item': item,
                                        'category': category
                                    }
                                    all_scopes.append(scope_info)
                            else:
                                logger.debug("No scope of work section found in document")
                    except Exception as e:
                        logger.error(f"Error processing file {file_path}: {str(e)}", exc_info=True)
    
    logger.info(f"Scope analysis complete. Found {len(all_scopes)} scope items")
    if all_scopes:
        results_df = pd.DataFrame(all_scopes)
    else:
        logger.warning("No scope items found in any documents")
        results_df = pd.DataFrame(columns=['project_path', 'source_file', 'scope_item', 'category'])
    
    return results_df

def categorize_scope_item(text: str) -> str:
    """
    Categorize a scope item based on keywords.
    This is a basic implementation - could be enhanced with ML/more sophisticated categorization.
    """
    text = text.lower()
    
    categories = {
        'ASSESSMENT': ['calculate', 'AIIC', 'ASTC', 'ASTM E336-17a', 'ASTM E1007-16'],
        'DESIGN': ['design', 'drawing', 'specification', 'detail'],
        'DOCUMENTATION': ['report', 'documentation', 'document', 'submittal'],
        'COORDINATION': ['coordination', 'meeting', 'review', 'consultation']
    }
    
    for category, keywords in categories.items():
        if any(keyword in text for keyword in keywords):
            return category
            
    return 'other'

def create_project_file_list(project_path: str | list[str], output_file: str = None) -> dict:
    """
    Create a text file listing all files in a project path and return a dictionary
    of folder structure with file counts.
    
    Args:
        project_path: Path to the project directory or list of project paths
        output_file: Path to output text file (if None, uses project name + '_files.txt')
        
    Returns:
        Dictionary with folder structure and file counts
    """
    # Handle both single path and list of paths
    if isinstance(project_path, list):
        all_folder_structures = {}
        for path in project_path:
            folder_structure = create_project_file_list(path)
            all_folder_structures.update(folder_structure)
        return all_folder_structures
    
    if output_file is None:
        # Create output filename based on project name
        project_name = os.path.basename(project_path)
        output_file = f"{project_name}_files.txt"
    
    # Dictionary to store folder structure and file counts
    folder_structure = {}
    
    # Open file for writing
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(f"File listing for project: {project_path}\n")
        f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("-" * 80 + "\n\n")
        
        # Walk through directory
        for root, dirs, files in os.walk(project_path):
            # Get relative path from project root
            rel_path = os.path.relpath(root, project_path)
            if rel_path == '.':
                rel_path = ''
            
            # Write folder path
            folder_indent = '  ' * (rel_path.count(os.sep))
            f.write(f"{folder_indent}📁 {os.path.basename(root)}/\n")
            
            # Count files in this folder
            file_count = len(files)
            
            # Add to structure dictionary
            folder_structure[rel_path] = file_count
            
            # Write files
            for file in files:
                file_indent = '  ' * (rel_path.count(os.sep) + 1)
                f.write(f"{file_indent}📄 {file}\n")
            
            f.write("\n")
    
    return folder_structure

def plot_folder_tree(project_path: str | list[str], folder_structure: dict = None):
    """Plot the folder tree visualization"""
    logger.info("Starting folder tree visualization")
    G = nx.DiGraph()
    cmap = get_cmap('viridis')

    if isinstance(project_path, list):
        logger.debug(f"Processing multiple project paths: {len(project_path)} projects")

        for path in project_path:
            logger.debug(f"Processing project path: {path}")
            if folder_structure is None:
                logger.debug("Generating folder structure")
                current_structure = create_project_file_list(path)
            else:
                current_structure = folder_structure
            project_name = os.path.basename(path)
            logger.debug(f"Adding root node: {project_name}")
            G.add_node(project_name, files=0)
            for folder_path, file_count in current_structure.items():
                logger.debug(f"Processing folder: {folder_path} with {file_count} files")
                if folder_path == '':
                    G.nodes[project_name]['files'] += file_count
                    continue
                components = folder_path.split(os.sep)
                parent = project_name
                for i, component in enumerate(components):
                    if i == 0:
                        node_path = component
                    else:
                        node_path = os.sep.join(components[:i+1])
                    if not G.has_node(node_path):
                        logger.debug(f"Adding new node: {node_path}")
                        G.add_node(node_path, files=0)
                        G.add_edge(parent, node_path)
                    if i == len(components) - 1:
                        G.nodes[node_path]['files'] += file_count
                    parent = node_path
        project_name = os.path.basename(project_path[0])
    else:
        logger.debug("Processing single project path")
        if folder_structure is None:
            logger.debug("Generating folder structure")
            current_structure = create_project_file_list(project_path)
        else:
            current_structure = folder_structure
        project_name = os.path.basename(project_path)
        logger.debug(f"Adding root node: {project_name}")
        G.add_node(project_name, files=0)
        for folder_path, file_count in current_structure.items():
            logger.debug(f"Processing folder: {folder_path} with {file_count} files")
            if folder_path == '':
                G.nodes[project_name]['files'] += file_count
                continue
            components = folder_path.split(os.sep)
            parent = project_name
            for i, component in enumerate(components):
                if i == 0:
                    node_path = component
                else:
                    node_path = os.sep.join(components[:i+1])
                if not G.has_node(node_path):
                    logger.debug(f"Adding new node: {node_path}")
                    G.add_node(node_path, files=0)
                    G.add_edge(parent, node_path)
                if i == len(components) - 1:
                    G.nodes[node_path]['files'] += file_count
                parent = node_path
    logger.debug("Calculating visualization parameters")
    max_files = max([data['files'] for _, data in G.nodes(data=True)]) if G.nodes else 1
    logger.debug(f"Maximum files in any folder: {max_files}")
    logger.debug("Creating plot")
    fig, ax = plt.subplots(figsize=(15, 10))
    pos = nx.spring_layout(G, k=1, iterations=50)
    logger.debug("Drawing network graph")
    nx.draw(G, pos,
            node_size=[2000 * (data['files'] / max_files) + 500 for _, data in G.nodes(data=True)],
            node_color=[cmap(data['files'] / max_files) if max_files > 0 else cmap(0) for _, data in G.nodes(data=True)],
            with_labels=True,
            labels={node: f"{node.split(os.sep)[-1]}\n({data['files']} files)" 
                   for node, data in G.nodes(data=True)},
            font_size=8,
            font_weight='bold',
            arrows=False,
            alpha=0.8,
            ax=ax)
    logger.debug("Adding colorbar and finalizing plot")
    sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(0, max_files))
    sm.set_array([])
    plt.colorbar(sm, ax=ax, label='Number of Files')
    ax.set_title(f"Folder Structure for {project_name}", fontsize=16)
    logger.info(f"Saving visualization to: {project_name}_folder_tree.png")
    plt.savefig(f"{project_name}_folder_tree.png", dpi=300, bbox_inches='tight')
    plt.close()
    logger.info("Folder tree visualization complete")

def format_file_tree(project_path: str) -> str:
    """
    Return a formatted string representing the file/folder tree for display.
    """
    lines = []
    for root, dirs, files in os.walk(project_path):
        rel_path = os.path.relpath(root, project_path)
        indent = '    ' * (rel_path.count(os.sep) if rel_path != '.' else 0)
        lines.append(f"{indent}📁 {os.path.basename(root)}/")
        for file in files:
            lines.append(f"{indent}    📄 {file}")
    return '\n'.join(lines)

def main():
    """Process multiple project folders to analyze proposal documents for dollar amounts."""
    # Example usage
    # root dir
    # //DLA-04/Shared/KAILUA PROJECTS/2023/
    project_paths = [
        "//DLA-04/Shared/KAILUA PROJECTS/2018/18-136 Masters at Ka'anapali Hillside Condos Floor Assessment",
        "//DLA-04/Shared/KAILUA PROJECTS/2018/18-099 Masters at Kaanpali AIIC Testing"
    ]
    
    # Analyze proposals and find dollar amounts
    # propsal_results = analyze_proposals(project_paths)
    
    # # Save results to CSV
    # propsal_results.to_csv('proposal_analysis_results.csv', index=False)
    
    # # Print summary
    # print(f"\nFound {len(propsal_results)} dollar amounts in proposal documents")
    # print("\nSample findings:")
    # print(propsal_results.head())

    # # Create project databases
    # project_dbs = [create_project_db(path) for path in project_paths]
    # print(project_dbs)
    # Save project databases to CSV
    # for folder_name, df in project_dbs.items():
    #     df.to_csv(f'{folder_name}_project_db.csv', index=False)

    # Get project paths for a specific year
    # year = '2023'
    # project_paths = get_year_project_paths("//DLA-04/Shared/KAILUA PROJECTS/", year)
    # print(project_paths)

    # # Analyze scope of work
    # scope_results = analyze_scope_of_work(project_paths)
    # print(scope_results)

    # Create project file list
    project_file_list = create_project_file_list(project_paths)
    print(project_file_list)

    # Plot folder tree
    plot_folder_tree(project_paths, project_file_list)

    # Format file tree
    formatted_tree = format_file_tree(project_paths[0])
    print(formatted_tree)

if __name__ == "__main__":
    main()
